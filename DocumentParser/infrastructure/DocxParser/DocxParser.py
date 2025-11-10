import re
from typing import List, Tuple

import docx
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from DocumentParser.domain.IDocumentParser import IDocumentParser
from DocumentParser.domain.ValueObjects.BoundingBox import BoundingBox
from DocumentParser.domain.ValueObjects.line import Line
from DocumentParser.domain.ValueObjects.word import Word
from DocumentParser.domain.entities.block import Block
from DocumentParser.domain.entities.document import Document
from DocumentParser.domain.entities.page import Page

from DocumentParser.infrastructure.DocxParser.DocxParserConstants import *
from DocumentParser.infrastructure.DocxParser import DocxParserValues


#TODO: Вынести в отдельный микросервис
class DocxParser(IDocumentParser):

    # TODO: Разбить на вспомогательные методы
    def process(self, path: str) -> Document:
        document = docx.Document(path)

        page_width_pt, page_height_pt = DocxParserValues.get_page_size(document)
        margins = DocxParserValues.get_page_margins(document)
        content_height = page_height_pt - margins['top'] - margins['bottom']

        pages = []
        current_page_blocks = []
        current_y = margins['top']
        current_page_start_y = margins['top']

        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                current_y += EMPTY_LINE_SPACING_PT
                if current_y - current_page_start_y > content_height:
                    if current_page_blocks:
                        pages.append(Page(blocks=current_page_blocks))
                    current_page_blocks = []
                    current_y = margins['top']
                    current_page_start_y = margins['top']
                continue

            paragraph_block = self._parse_paragraph(paragraph, page_width_pt, margins, current_y)
            if not paragraph_block.lines:
                continue

            paragraph_height = DocxParserValues.get_paragraph_height(paragraph_block, current_y) - current_y

            if current_y + paragraph_height - current_page_start_y > content_height:
                if current_page_blocks:
                    pages.append(Page(blocks=current_page_blocks))

                current_page_blocks = []
                current_y = margins['top']
                current_page_start_y = margins['top']

                paragraph_block = self._parse_paragraph(paragraph, page_width_pt, margins, current_y)
                if not paragraph_block.lines:
                    continue

            current_page_blocks.append(paragraph_block)
            current_y = DocxParserValues.get_paragraph_height(paragraph_block, current_y)

            if current_y - current_page_start_y > content_height:
                pages.append(Page(blocks=current_page_blocks))
                current_page_blocks = []
                current_y = margins['top']
                current_page_start_y = margins['top']

        if current_page_blocks:
            pages.append(Page(blocks=current_page_blocks))

        if not pages:
            pages.append(Page(blocks=[]))

        return Document(
            path=path,
            source_type='docx',
            pages=pages
        )

    def _parse_paragraph(self, paragraph: Paragraph, page_width: float, margins: dict, start_y: float) -> Block:
        words_with_styles = self._get_words_with_styles(paragraph)
        if not words_with_styles:
            return Block(lines=[])

        left_indent = DocxParserValues.get_left_paragraph_indent(paragraph)
        right_indent = DocxParserValues.get_right_paragraph_indent(paragraph)
        available_width = page_width - margins['left'] - margins['right'] - left_indent - right_indent

        alignment = paragraph.paragraph_format.alignment
        line_start_x_offset = margins['left'] + left_indent

        lines = self._layout_lines_in_paragraph(
            words_with_styles,
            start_y,
            available_width,
            alignment,
            line_start_x_offset
        )

        return Block(lines=lines)

    def _layout_lines_in_paragraph(
            self,
            words_with_styles: List[Tuple],
            start_y: float,
            available_width: float,
            alignment: WD_PARAGRAPH_ALIGNMENT,
            left_offset: float) -> List[Line]:
        lines = []
        current_line_y = start_y
        current_line_words = []
        current_line_width = 0

        for word_text, font_size, is_bold in words_with_styles:
            word_width = DocxParserValues.get_word_width(word_text, font_size, is_bold)

            if current_line_words and current_line_width + word_width > available_width:
                line = self._get_line_from_words(
                    current_line_words,
                    current_line_y,
                    available_width,
                    alignment,
                    left_offset
                )
                lines.append(line)

                current_line_y += DocxParserValues.get_line_height(current_line_words)
                current_line_words = []
                current_line_width = 0

            current_line_words.append((word_text, font_size, is_bold, word_width))
            current_line_width += word_width

        if current_line_words:
            line = self._get_line_from_words(
                current_line_words,
                current_line_y,
                available_width,
                alignment,
                left_offset
            )
            lines.append(line)

        return lines

    # TODO: Отказаться от Tuple в пользу собственной структуры данных
    def _get_words_with_styles(self, paragraph: Paragraph) -> List[Tuple[str, float, bool]]:
        words = []
        for run in paragraph.runs:
            font_size = DocxParserValues.get_font_size(run)
            is_bold = run.font.bold or False
            run_words = re.findall(r'\S+', run.text)

            for word in run_words:
                words.append((word, font_size, is_bold))
        return words

    def _get_line_from_words(
            self,
            words: List[Tuple],
            y: float,
            available_width: float,
            alignment,
            left_offset: float
    ) -> Line:
        line_words = []
        total_width = sum(word[3] for word in words)
        current_x = DocxParserValues.get_line_start_x(alignment, left_offset, available_width, total_width)

        for word_text, font_size, is_bold, word_width in words:
            word_height = font_size * LINE_HEIGHT_FACTOR
            bounding_box = BoundingBox(
                left=current_x,
                top=y,
                right=current_x + word_width,
                bottom=y + word_height
            )
            line_words.append(Word(word=word_text, bounding_box=bounding_box))
            current_x += word_width

        return Line(words=line_words)